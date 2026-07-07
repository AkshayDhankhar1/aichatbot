"""Read-only document loading for the ingestion pipeline.

Handles the formats actually present in `docs/`: CSV, TSV, JSON, XLSX, DOCX,
PDF, Markdown and plain text. Everything here READS ONLY — it never writes,
edits, or deletes anything under `docs/`.

Each loaded item is a small dataclass carrying the raw text plus the source
filename (used later for chunk metadata and inline source attribution).
"""
from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path

# Directories / files to always ignore (macOS zip cruft, hidden files).
IGNORED_DIR_NAMES = {"__MACOSX", ".git", "node_modules"}
IGNORED_FILE_PREFIXES = ("._", ".")  # AppleDouble files, dotfiles


@dataclass
class LoadedDoc:
    source: str   # filename, e.g. "deals.csv" — becomes the [source] citation tag
    text: str     # extracted plain-text content


def _iter_files(docs_dir: Path):
    """Yield candidate files under docs_dir, skipping junk."""
    for path in sorted(docs_dir.rglob("*")):
        if not path.is_file():
            continue
        if any(part in IGNORED_DIR_NAMES for part in path.parts):
            continue
        if path.name.startswith(IGNORED_FILE_PREFIXES):
            continue
        yield path


# --------------------------------------------------------------------------- #
# Per-format extractors. Each returns plain text (best-effort).                #
# --------------------------------------------------------------------------- #
def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_delimited(path: Path, delimiter: str) -> str:
    """CSV/TSV -> readable 'col: value' rows so a chunk stays self-describing."""
    rows_out: list[str] = []
    with path.open(newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f, delimiter=delimiter)
        rows = list(reader)
    if not rows:
        return ""
    header = rows[0]
    for i, row in enumerate(rows[1:], start=1):
        pairs = [f"{h}: {v}" for h, v in zip(header, row)]
        rows_out.append(f"Row {i} | " + " | ".join(pairs))
    return f"Table {path.name} (columns: {', '.join(header)})\n" + "\n".join(rows_out)


def _load_json(path: Path) -> str:
    data = json.loads(_load_text(path))

    def flatten(obj, prefix=""):
        lines = []
        if isinstance(obj, dict):
            parts = [f"{k}: {v}" for k, v in obj.items() if not isinstance(v, (dict, list))]
            if parts:
                lines.append((prefix + " | " if prefix else "") + " | ".join(parts))
            for k, v in obj.items():
                if isinstance(v, (dict, list)):
                    lines.extend(flatten(v, f"{prefix}{k}."))
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                lines.extend(flatten(item, f"{prefix}[{i}]"))
        else:
            lines.append(f"{prefix}{obj}")
        return lines

    return f"JSON {path.name}\n" + "\n".join(flatten(data))


def _load_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    out: list[str] = []
    for ws in wb.worksheets:
        out.append(f"Sheet: {ws.title}")
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = [str(c) if c is not None else "" for c in rows[0]]
        for i, row in enumerate(rows[1:], start=1):
            pairs = [f"{h}: {v}" for h, v in zip(header, row) if v is not None]
            if pairs:
                out.append(f"Row {i} | " + " | ".join(pairs))
    wb.close()
    return f"Workbook {path.name}\n" + "\n".join(out)


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # Include table cell text too.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return f"Document {path.name}\n" + "\n".join(paras)


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        if txt.strip():
            pages.append(txt)
    return f"PDF {path.name}\n" + "\n".join(pages)


_EXTRACTORS = {
    ".txt": _load_text,
    ".md": _load_text,
    ".markdown": _load_text,
    ".csv": lambda p: _load_delimited(p, ","),
    ".tsv": lambda p: _load_delimited(p, "\t"),
    ".json": _load_json,
    ".xlsx": _load_xlsx,
    ".xlsm": _load_xlsx,
    ".docx": _load_docx,
    ".pdf": _load_pdf,
}


def load_docs(docs_dir: Path) -> list[LoadedDoc]:
    """Load every supported file under docs_dir. Read-only."""
    docs: list[LoadedDoc] = []
    for path in _iter_files(docs_dir):
        extractor = _EXTRACTORS.get(path.suffix.lower())
        if extractor is None:
            print(f"  · skipping unsupported file: {path.name}")
            continue
        try:
            text = extractor(path)
        except Exception as exc:  # keep going even if one file is malformed
            print(f"  ! failed to load {path.name}: {exc}")
            continue
        if text and text.strip():
            docs.append(LoadedDoc(source=path.name, text=text))
            print(f"  · loaded {path.name} ({len(text)} chars)")
    return docs
